# Imports Section
from docx import Document
import pandas as pd
from io import BytesIO
import streamlit as st
from docx2pdf import convert
import tempfile
import os


# List of registered doctors with their names and professional numbers
registered_doctors = [
    {
        "doctor_name": "Antonio Buzido Jimenez",
        "number": "Médico colegiado Nº 9800 de Sevilla",
    },
    {
        "doctor_name": "Dra. Paz Marian Casal",
        "number": "Médico colegiado Nº 987 de Sevilla",
    },
]

# Auxiliar Functions
def generate_pdf_with_docx2pdf(doc):
    import pythoncom
    pythoncom.CoInitialize()  # Manually initialize COM

    with tempfile.TemporaryDirectory() as tmpdir:
        docx_path = os.path.join(tmpdir, "temp.docx")
        pdf_path = os.path.join(tmpdir, "temp.pdf")

        doc.save(docx_path)
        convert(docx_path, pdf_path)

        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()

    return pdf_bytes

# Core Functions - - - -
@st.dialog("Generar Informe Médico", width="large")
def preview_file(final_df, default_open_status):
    """
    Displays a dialog for generating medical reports using a .docx template.

    This dialog allows the user to:
    - Upload a template (.docx) file.
    - Select a doctor from a list of registered doctors.
    - Choose a column from a DataFrame to be extracted into the report.
    - Provide additional details such as case number, documentation given, and documentation not given.

    Args:
        final_df (pd.DataFrame): The DataFrame containing the data that will populate the report.
    """

    # Simple toggle using session_state
    st.session_state.show_download = default_open_status
    
    col1, col2 = st.columns([2, 2])  # Adjust width ratio for the columns

    # File uploader for the user to select a Word template
    template_file = st.file_uploader(
        "Por favor selecciona la plantilla deseada", type="docx"
    )

    max_col = final_df.shape[1] - 1  # Maximum column index
    column_indices = list(range(0, max_col + 1))  # List of available column indices

    # Create a list of doctor names for the dropdown
    doctor_names = [doc["doctor_name"] for doc in registered_doctors]

    # with col1:
    #     # Selectbox for choosing the doctor
    #     doctor = st.selectbox("Doctor", doctor_names)
    # with col2:
    #     # Selectbox for choosing the column to extract from the DataFrame
    #     selected_column = st.selectbox("Columna a extraer", column_indices)

    # # Text inputs for additional information required in the report
    # expedient_number = st.text_input("Número de Expediente")
    # documentation_given = st.text_input("Documentación Aportada")
    # documentation_not_given = st.text_input("Documentación no Aportada")

    # Initialize session state for tracking changes
    if "doctor" not in st.session_state:
        st.session_state.doctor = ""
    if "selected_column" not in st.session_state:
        st.session_state.selected_column = ""
    if "expedient_number" not in st.session_state:
        st.session_state.expedient_number = ""
    if "documentation_given" not in st.session_state:
        st.session_state.documentation_given = ""
    if "documentation_not_given" not in st.session_state:
        st.session_state.documentation_not_given = ""

    # Selectbox for choosing the doctor
    with col1:
        doctor = st.selectbox("Doctor", doctor_names)
    # Selectbox for choosing the column to extract from the DataFrame
    with col2:
        selected_column = st.selectbox("Columna a extraer", column_indices)
        
    # Text inputs for additional information required in the report
    expedient_number = st.text_input("Número de Expediente")
    documentation_given = st.text_input("Documentación Aportada")
    documentation_not_given = st.text_input("Documentación no Aportada")

    # Check if any input has changed
    if doctor != st.session_state.doctor:
        st.session_state.show_download = False
        st.session_state.doctor = doctor

    if selected_column != st.session_state.selected_column:
        st.session_state.show_download = False
        st.session_state.selected_column = selected_column

    if expedient_number != st.session_state.expedient_number:
        st.session_state.show_download = False
        st.session_state.expedient_number = expedient_number

    if documentation_given != st.session_state.documentation_given:
        st.session_state.show_download = False
        st.session_state.documentation_given = documentation_given

    if documentation_not_given != st.session_state.documentation_not_given:
        st.session_state.show_download = False
        st.session_state.documentation_not_given = documentation_not_given


    # Find the selected doctor in the list of registered doctors
    matched_doctor = next(
        (doc for doc in registered_doctors if doc["doctor_name"] == doctor), "N/A"
    )

    # Collect the additional information to be used in the document
    additional_documentation = {
        "{{Doctor}}": doctor,
        "{{Numero de colegiado}}": matched_doctor["number"],
        "{{Doctor Identification}}": matched_doctor["number"],
        "{{Expediente}}": expedient_number,
        "{{Documentación aportada}}": documentation_given,
        "{{Documentación no aportada}}": documentation_not_given,
    }


    
    if st.button("Procesar información", use_container_width=True):
        st.session_state.show_download = True
    
    # Trigger document generation and offer download options if the template is uploaded
    if st.session_state.show_download == True:
        if template_file:
            fill_and_offer_multiple_downloads(
                final_df, selected_column, template_file, additional_documentation
            )


def fill_and_offer_multiple_downloads(
    df: pd.DataFrame, column_index: int, template_path: str, extra_information: dict
):
    """
    Fills a Word template with data from a given DataFrame column and shows download buttons for both DOCX and PDF formats.

    This function:
    - Extracts the data from the DataFrame column.
    - Replaces placeholders in the Word document with the extracted data and additional information.
    - Provides download buttons for both DOCX and PDF formats.

    Args:
        df (pd.DataFrame): The DataFrame containing data to be inserted into the report.
        column_index (int): The index of the column in the DataFrame to extract.
        template_path (str): Path to the .docx template that will be used to generate the report.
        extra_information (dict): A dictionary of additional information to be inserted into the template.
    """

    # Layout: two columns for the download buttons
    col1, col2 = st.columns(2)

    # Extract the column data to be used in the template
    column_data = df[column_index]

    # Prepare the replacements for placeholders in the template
    replacements = {
        f"{{{{{key}}}}}": (
            (str(value).split(" Hora: ")[0] if key == "Fecha siniestro" else str(value))
            if pd.notnull(value)
            else ""
        )
        for key, value in column_data.items()
    }

    # Special handling for the 'Fecha siniestro' field to separate date and time
    if "Fecha siniestro" in column_data and pd.notnull(column_data["Fecha siniestro"]):
        fecha_value = column_data["Fecha siniestro"]
        if " Hora: " in fecha_value:
            replacements["{{Fecha Siniestro}}"] = fecha_value.split(" Hora: ")[0]
            replacements["{{Hora}}"] = fecha_value.split(" Hora: ")[1]

    # Load the template Word document
    doc = Document(template_path)

    # Replace the placeholders in the paragraphs of the document
    for para in doc.paragraphs:
        for key, val in replacements.items():
            if key in para.text:
                para.text = para.text.replace(key, val)
        for key, val in extra_information.items():
            if key in para.text:
                para.text = para.text.replace(key, val)

    # Replace placeholders in the tables of the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)
                for key, val in extra_information.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, val)

    # Save the modified document to an in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)


    pdf_bytes = generate_pdf_with_docx2pdf(doc)
    if pdf_bytes != "":
        # Download button for DOCX format
        with col1:
            file_name = (
                f"informe_col{column_index}.docx"  # Dynamic file name based on column index
            )
            st.download_button(
                label=f"📄 Descargar Informe en formato .docx",
                data=buffer,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )

        # Generate PDF
        with col2:
            
            file_name = f"informe_col{column_index}.pdf"  # Dynamic file name based on column index
            st.download_button(
                label="📄 Descargar Informe en formato .pdf",
                data=pdf_bytes,
                file_name=file_name,
                mime="application/pdf",
                use_container_width=True,
            )
    else:
        st.title("Preparado informe")
